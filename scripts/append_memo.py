"""
Appends the Act V Executive Decision Memo sections to Estifanos_Interim_report.docx.
Run from the repo root: python scripts/append_memo.py
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCX_PATH = "Estifanos_Interim_report.docx"


def add_heading(doc, text, level=1):
    if level <= 2:
        p = doc.add_heading(text, level=level)
    else:
        # Document only has H1/H2; emulate H3 as bold Normal
        p = doc.add_paragraph(style="Normal")
        run = p.add_run(text)
        run.bold = True
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Paragraph")
    p.add_run(f"•  {text}")
    return p


def add_table_row(table, cells):
    row = table.add_row()
    for i, val in enumerate(cells):
        row.cells[i].text = val
    return row


def build_memo(doc):
    doc.add_page_break()

    # ── Title block ──────────────────────────────────────────────────────────
    t = doc.add_heading("Act V — Executive Decision Memo", level=1)
    doc.add_paragraph("To: CEO / CFO, Tenacious Consulting & Outsourcing")
    doc.add_paragraph("From: Estifanos Teklay Amare")
    doc.add_paragraph("Date: 25 April 2026")
    doc.add_paragraph("Re: SignalForge — Pilot Recommendation and Honest Performance Assessment")
    doc.add_paragraph("")

    # ── 1. Executive Summary ─────────────────────────────────────────────────
    add_heading(doc, "1. Executive Summary", level=2)
    add_body(
        doc,
        "SignalForge is a five-agent outbound system that converts public hiring, funding, "
        "and leadership signals into personalised cold emails, reply-thread management, and "
        "Cal.com discovery-call bookings — achieving a conversational task pass@1 of 51.3% "
        "(95% CI [43.3%, 59.3%]) against the tau²-Bench retail baseline, which exceeds "
        "the published voice-agent ceiling of ~42% for this task family. "
        "Signal-grounded outreach is benchmarked at 7–12% reply rate versus 1–3% "
        "for generic cold email, a delta of +6–9 percentage points that is consistent "
        "with Clay 2025 and Smartlead 2025 published case studies. "
        "The recommendation is to run a 30-day Segment 2 (cost-restructuring) pilot at "
        "150 leads per week with a ≈18-USD weekly all-in budget and a success criterion "
        "of ≥10% reply rate and ≥20 discovery calls booked by Day 30."
    )

    # ── 2. Cost per Qualified Lead ───────────────────────────────────────────
    add_heading(doc, "2. Cost per Qualified Lead", level=2)
    add_body(
        doc,
        "A 'qualified lead' is a prospect who has replied to at least one outbound email "
        "AND answered at least one segment-specific qualification question from the "
        "ConversationAgent, entering the QUALIFIED state in the pipeline state machine."
    )
    add_body(doc, "Input decomposition (dev tier, 60 prospects/week):")

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Normal Table"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Cost Item"
    hdr[1].text = "Basis"
    hdr[2].text = "Weekly Cost"
    add_table_row(tbl, [
        "LLM (deepseek-v3 via OpenRouter)",
        "3 agent calls × ~9,000 tokens avg × $0.68/M blended = $0.004/prospect × 60",
        "$0.24"
    ])
    add_table_row(tbl, [
        "Compute / rig",
        "Cloud VM + Playwright (estimated at dev scale)",
        "~$5.00"
    ])
    add_table_row(tbl, [
        "APIs (Resend free tier, HubSpot sandbox, Cal.com free)",
        "Free tiers cover dev volume",
        "$0.00"
    ])
    add_table_row(tbl, [
        "Africa's Talking SMS",
        "~10 warm-lead SMS/week × $0.015",
        "$0.15"
    ])
    add_table_row(tbl, ["Total", "", "~$5.39 / week"])

    doc.add_paragraph("")
    add_body(doc, "Qualified leads per week derivation:")
    add_bullet(doc, "60 prospects × 9.5% reply rate (midpoint 7–12%) = 5.7 replies")
    add_bullet(doc, "5.7 replies × 55% qualification rate = 3.1 qualified leads")

    add_body(
        doc,
        "Cost per qualified lead = $5.39 / 3.1 = $1.74 (dev tier). "
        "At eval tier (claude-sonnet-4-6, ~7× LLM cost): $1.68 LLM/week → CPL ≈ $2.21. "
        "Both figures are well within the cost envelope documented in seed/baseline_numbers.md."
    )

    # ── 3. Stalled-Thread Rate Delta ─────────────────────────────────────────
    add_heading(doc, "3. Stalled-Thread Rate Delta", level=2)
    add_body(
        doc,
        "Definition: a thread is 'stalled' if the ConversationAgent produces no outbound action "
        "within 72 hours of receiving an inbound reply. This covers the reply-to-qualification "
        "stage and is distinct from the late-stage deal stall (72% per seed/baseline_numbers.md), "
        "which reflects deals already in CRM pipeline and is outside the scope of this system."
    )

    tbl2 = doc.add_table(rows=1, cols=3)
    tbl2.style = "Normal Table"
    hdr2 = tbl2.rows[0].cells
    hdr2[0].text = "Process"
    hdr2[1].text = "Stall Rate"
    hdr2[2].text = "Source"
    add_table_row(tbl2, [
        "Manual SDR (reply → qualification)",
        "30–40%",
        "LeadIQ / Apollo 2026 benchmarks"
    ])
    add_table_row(tbl2, [
        "SignalForge automated (webhook-driven)",
        "<1% (responds within seconds of reply webhook)",
        "Measured in dev environment; all replies handled"
    ])
    add_table_row(tbl2, ["Delta", "−30 to −40 percentage points", ""])

    doc.add_paragraph("")
    add_body(
        doc,
        "Caveat: these measurements use synthetic prospects in a development environment "
        "(OUTBOUND_ENABLED=false). Transfer to production may introduce latency spikes "
        "from webhook delivery failures or API rate limits not observed at dev scale. "
        "The late-stage stall rate of 72% (seed/baseline_numbers.md) remains unchanged by "
        "this system, as it governs deal-stage dynamics after the discovery call is booked."
    )

    # ── 4. Competitive-Gap Reply Rate Delta ──────────────────────────────────
    add_heading(doc, "4. Competitive-Gap Outbound Reply-Rate Delta", level=2)
    add_body(
        doc,
        "Two outbound variants are defined in the pipeline via the ICP confidence gate:"
    )
    add_bullet(
        doc,
        "Variant A — Signal-grounded: email grounded in a competitor_gap_brief, "
        "ICP-specific pitch language, and public hiring signal evidence. Applied when "
        "confidence_score ≥ 0.6 and segment is assigned."
    )
    add_bullet(
        doc,
        "Variant B — Generic exploratory: no competitor gap brief, no segment-specific "
        "pitch language. Applied when confidence_score < 0.6 (abstain path)."
    )

    tbl3 = doc.add_table(rows=1, cols=4)
    tbl3.style = "Normal Table"
    h3 = tbl3.rows[0].cells
    h3[0].text = "Variant"
    h3[1].text = "Reply Rate"
    h3[2].text = "Source / Sample"
    h3[3].text = "Notes"
    add_table_row(tbl3, [
        "A — Signal-grounded",
        "7–12% (midpoint 9.5%)",
        "Clay 2025, Smartlead 2025 case studies (seed/baseline_numbers.md)",
        "Top-quartile signal-grounded outbound benchmark"
    ])
    add_table_row(tbl3, [
        "B — Generic cold",
        "1–3% (midpoint 2%)",
        "LeadIQ 2026, Apollo 2026 benchmarks (seed/baseline_numbers.md)",
        "B2B cold-email industry baseline"
    ])
    add_table_row(tbl3, [
        "Delta",
        "+6 to +9 percentage points",
        "",
        "4.8× relative at midpoints"
    ])

    doc.add_paragraph("")
    add_body(
        doc,
        "Limitation: this comparison is based on published industry benchmarks used to "
        "justify the signal-grounded design, not a live A/B test on Tenacious-specific "
        "prospects. OUTBOUND_ENABLED=false in development means no real emails have been "
        "sent. The delta should be treated as directional. A 30-day pilot (Section 6) "
        "would produce the first live measurement."
    )

    # ── 5. AI Maturity Scoring Limitations ──────────────────────────────────
    add_heading(doc, "5. Public-Signal Lossiness of AI Maturity Scoring", level=2)
    add_body(
        doc,
        "The AI maturity score (0–3) is computed from five public signals. "
        "Two systematic error modes are known:"
    )

    add_heading(doc, "False Positive Mode (scores high, should score low)", level=3)
    add_body(
        doc,
        "Archetype: an AI talent platform or staffing agency (e.g., a company whose "
        "product is matching AI engineers to clients). These firms have a high fraction "
        "of AI-adjacent job postings because they are recruiting for client placements, "
        "not building internal AI systems. They frequently have named 'AI Lead' or "
        "'ML Director' roles posted permanently."
    )
    add_body(
        doc,
        "Agent action: scores 2–3 → classified as Segment 4 capability-gap prospect "
        "→ InsightAgent generates a competitor gap brief framing them as lagging "
        "top-quartile peers in AI adoption."
    )
    add_body(
        doc,
        "Business impact: the prospect is either a competitor or a company that sells AI "
        "services; a Segment 4 gap brief is actively insulting to them. The thread is closed "
        "immediately, and the company may flag Tenacious as a low-quality sender to their "
        "network. The explicit constraint in seed/icp_definition.md ('Segment 4 pitch to a "
        "score-0 prospect damages brand') applies symmetrically to this case."
    )

    add_heading(doc, "False Negative Mode (scores low, should score high)", level=3)
    add_body(
        doc,
        "Archetype: a deep-tech AI startup with a fully hired founding AI team. They have "
        "no open AI job postings (the team is already built), a closed-source model (no "
        "public GitHub activity), and minimal press coverage (stealth or pre-launch). "
        "GitHub activity signal is currently unimplemented (returns None; see probe PL-032)."
    )
    add_body(
        doc,
        "Agent action: scores 0 → abstains → sends only a generic exploratory "
        "email with no competitor gap brief. The gap brief is never generated."
    )
    add_body(
        doc,
        "Business impact: the company that would benefit most from a Tenacious "
        "capability-extension pitch (Segment 4 at full confidence) is never pitched "
        "specifically. The system under-reaches the highest-value sub-population of "
        "Segment 4 prospects systematically. Implementing the GitHub activity signal "
        "(currently marked None in signal_computer.py) would partially address this."
    )

    # ── 6. Pilot Recommendation ──────────────────────────────────────────────
    add_heading(doc, "6. Pilot Recommendation: Segment 2, 30-Day Pilot", level=2)
    add_body(
        doc,
        "Segment justification: Segment 2 (mid-market cost restructuring) is recommended "
        "over Segments 1, 3, and 4 for the initial pilot because: (1) the layoffs.fyi "
        "signal is the highest-quality and most verifiable of the four signal categories "
        "— layoff events are public, timestamped, and not subject to the AI-maturity "
        "scoring false-positive modes described above; (2) the buying moment is well-defined "
        "(120-day window post-layoff with active engineering hiring); "
        "(3) the cost-lever pitch angle has the strongest benchmark support in the "
        "discovery-to-proposal conversion data (30–50%, seed/baseline_numbers.md)."
    )

    tbl4 = doc.add_table(rows=1, cols=2)
    tbl4.style = "Normal Table"
    h4 = tbl4.rows[0].cells
    h4[0].text = "Pilot Parameter"
    h4[1].text = "Value"
    add_table_row(tbl4, ["Segment", "Segment 2 — cost restructuring"])
    add_table_row(tbl4, ["Duration", "30 days"])
    add_table_row(tbl4, ["Lead volume", "150 leads per week (2.5× dev-scale SDR target)"])
    add_table_row(tbl4, [
        "Weekly budget",
        "~$18/week all-in (LLM: $0.60, rig: $12, APIs: $0.38, buffer: $5)"
    ])
    add_table_row(tbl4, ["Total 30-day budget", "~$72"])
    add_table_row(tbl4, [
        "Success criterion",
        (
            "≥10% reply rate on signal-grounded Segment 2 outreach (vs 1–3% cold baseline) "
            "AND ≥20 discovery calls booked by Day 30, measured from HubSpot "
            "activity log and Langfuse trace count"
        )
    ])
    add_table_row(tbl4, [
        "Go / No-Go decision point",
        "Day 15 interim check: if reply rate < 5% across ≥300 delivered emails, "
        "pause and review signal quality before Day 30"
    ])

    # ── 7. Honest Unresolved Failure ─────────────────────────────────────────
    add_heading(doc, "7. Honest Unresolved Failure: Gap Over-Claiming", level=2)
    add_body(
        doc,
        "The mechanism designed in method.md (multi-stage bench validation gate) addresses "
        "bench over-commitment (Category 3, trigger rate 11.3%). It does not address "
        "gap over-claiming (Category 10, trigger rate 21.7%)."
    )
    add_body(
        doc,
        "Specific failure (probe PL-033): InsightAgent's LLM call infers a competitive trend "
        "(e.g., 'all three of your top competitors have launched dedicated AI platform teams') "
        "from one or two weak peer evidence points. The gap_quality_self_check field is "
        "present in the output schema but its value is not enforced as a blocking condition "
        "by the calling code — a self-check of 'low' does not trigger a retry."
    )
    add_body(
        doc,
        "Triggering conditions: InsightAgent receives fewer than 5 viable peers from "
        "load_companies_by_industry(), or the peers returned have no recent AI signals "
        "(all scores 0). The LLM generates gap findings from the partial data anyway, "
        "stating trends as if fully evidenced."
    )
    add_body(
        doc,
        "Business impact if deployed: at a 21.7% trigger rate, approximately 1 in 5 Segment 4 "
        "gap briefs contains an unsupported competitive claim. Segment 4's primary buyer is a "
        "CTO with deep domain knowledge who will immediately recognise when a competitive "
        "intelligence claim is inflated or unsourced. Detection by an expert buyer permanently "
        "closes the thread and produces a negative brand signal. "
        "At 150 leads/week with ~40% qualifying as Segment 4, approximately 13 Segment 4 "
        "leads per week generate a gap brief; at 21.7% trigger rate, ~3 contain an unsupported "
        "claim per week. Even at conservative detection rates, this is a material source of "
        "wasted Segment 4 pipeline per month."
    )
    add_body(
        doc,
        "Path to resolution (not implemented in this submission): enforce gap_quality_self_check "
        "as a BLOCK condition in GuardrailAgent, requiring InsightAgent to retry with a "
        "'no gap finding' fallback when evidence count < 2 per finding. This is a one-sprint "
        "fix with a clear acceptance test (PL-033 trigger rate drops to < 5%)."
    )


def main():
    if not os.path.exists(DOCX_PATH):
        print(f"ERROR: {DOCX_PATH} not found. Run from repo root.")
        return

    doc = Document(DOCX_PATH)
    build_memo(doc)
    doc.save(DOCX_PATH)
    print(f"Memo sections appended to {DOCX_PATH} successfully.")


if __name__ == "__main__":
    main()
